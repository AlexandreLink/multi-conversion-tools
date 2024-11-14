import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document

st.title("Fusion de Noms pour Remerciements")

# Téléversement des deux fichiers Excel, acceptant à la fois .xls et .xlsx
file_rem = st.file_uploader("Téléversez le fichier de remerciements (Excel)", type=["xls", "xlsx"])
file_changes = st.file_uploader("Téléversez le fichier de changements de noms (Excel)", type=["xls", "xlsx"])

if file_rem and file_changes:
    # Charger les deux fichiers
    df_rem = pd.read_excel(file_rem)
    df_changes = pd.read_excel(file_changes)

    # Vérifier que les colonnes nécessaires sont présentes
    if 'Reference' not in df_rem.columns or 'Commande' not in df_changes.columns:
        st.error("Le fichier de remerciements doit contenir la colonne 'Reference' et le fichier de changements doit contenir la colonne 'Commande'.")
    elif 'A supprimer des pages Remerciements' not in df_changes.columns or 'A faire apparaitre sur les pages Remerciements' not in df_changes.columns:
        st.error("Le fichier de changements doit contenir les colonnes 'A supprimer des pages Remerciements' et 'A faire apparaitre sur les pages Remerciements'.")
    else:
        # Préparer les références de commandes en supprimant le symbole "#" dans df_changes
        df_changes['Commande'] = df_changes['Commande'].str.replace('#', '')

        # Créer un dictionnaire pour les remplacements basés sur les références de commandes
        replacements = df_changes.set_index('Commande')['A faire apparaitre sur les pages Remerciements'].to_dict()

        # Appliquer les remplacements
        df_rem['Nom Complet'] = df_rem.apply(
            lambda row: replacements.get(row['Reference'], f"{row['Prénom']} {row['Nom']}"), axis=1
        )

        # Trier les noms par ordre alphabétique
        sorted_names = sorted(df_rem['Nom Complet'].unique())

        # Créer un document Word et ajouter les noms
        doc = Document()
        doc.add_heading("Remerciements", 0)
        doc.add_paragraph(", ".join(sorted_names))

        # Enregistrer le fichier Word dans un flux de mémoire pour le téléchargement
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Bouton de téléchargement pour le fichier Word
        st.download_button(
            label="Télécharger le fichier de remerciements (Word)",
            data=buffer,
            file_name="Remerciements.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
