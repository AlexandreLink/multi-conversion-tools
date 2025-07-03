import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
import unicodedata

def detect_export_type(df):
    """
    Détecte automatiquement le type d'export basé sur les colonnes présentes
    Returns: 'samourai' ou 'pirate'
    """
    columns = df.columns.tolist()
    
    # Type Samouraï : colonnes Reference, Nom, Prénom
    if 'Reference' in columns and 'Nom' in columns and 'Prénom' in columns:
        return 'samourai'
    
    # Type Pirate : colonne # et Nom complet ou colonnes de facturation
    elif '#' in columns and ('Nom complet' in columns or 
                            ('Prénom de facturation' in columns and 'Nom de facturation' in columns)):
        return 'pirate'
    
    else:
        return 'unknown'

def standardize_dataframe(df, export_type):
    """
    Standardise le DataFrame selon le type d'export détecté
    Returns: DataFrame avec colonnes Reference, Nom, Prénom
    """
    if export_type == 'samourai':
        # Déjà au bon format, juste nettoyer les espaces
        df['Reference'] = df['Reference'].astype(str).str.strip()
        df['Nom'] = df['Nom'].astype(str).str.strip()
        df['Prénom'] = df['Prénom'].astype(str).str.strip()
        return df
    
    elif export_type == 'pirate':
        # Créer un nouveau DataFrame standardisé
        standardized_df = pd.DataFrame()
        
        # Utiliser la colonne # comme Reference
        standardized_df['Reference'] = df['#'].astype(str).str.strip()
        
        # Fonction pour obtenir le meilleur nom/prénom disponible
        def get_best_name(row):
            # Priorité 1: Nom et prénom de facturation (vrais noms)
            prenom_fact = row.get('Prénom de facturation', '')
            nom_fact = row.get('Nom de facturation', '')
            
            if (pd.notna(prenom_fact) and pd.notna(nom_fact) and 
                str(prenom_fact).strip() and str(nom_fact).strip()):
                return str(prenom_fact).strip(), str(nom_fact).strip()
            
            # Priorité 2: Nom et prénom de livraison (vrais noms)
            prenom_livr = row.get('Prénom de livraison', '')
            nom_livr = row.get('Nom de livraison', '')
            
            if (pd.notna(prenom_livr) and pd.notna(nom_livr) and 
                str(prenom_livr).strip() and str(nom_livr).strip()):
                return str(prenom_livr).strip(), str(nom_livr).strip()
            
            # Priorité 3: Nom complet (que ce soit un vrai nom ou un pseudo)
            nom_complet = row.get('Nom complet', '')
            if pd.notna(nom_complet) and str(nom_complet).strip():
                nom_complet = str(nom_complet).strip()
                parts = nom_complet.split()
                
                if len(parts) >= 2:
                    # Premier mot = prénom, le reste = nom
                    prenom = parts[0]
                    nom = ' '.join(parts[1:])
                    return prenom, nom
                elif len(parts) == 1:
                    return parts[0], ''
            
            # Priorité 4: Utiliser ce qui est disponible même si incomplet
            if pd.notna(prenom_fact) and str(prenom_fact).strip():
                return str(prenom_fact).strip(), str(nom_fact).strip() if pd.notna(nom_fact) else ''
            if pd.notna(prenom_livr) and str(prenom_livr).strip():
                return str(prenom_livr).strip(), str(nom_livr).strip() if pd.notna(nom_livr) else ''
            
            return '', ''
        
        # Appliquer la logique à chaque ligne
        names_data = df.apply(get_best_name, axis=1)
        standardized_df['Prénom'] = [x[0] for x in names_data]
        standardized_df['Nom'] = [x[1] for x in names_data]
        
        # Copier les autres colonnes importantes si elles existent
        for col in ['Email', 'E-mail', 'Pseudo', 'Identifiant / Pseudonyme']:
            if col in df.columns:
                standardized_df[col] = df[col]
                break
        
        return standardized_df
    
    else:
        raise ValueError("Type d'export non reconnu")

def get_export_info(export_type, df):
    """
    Retourne les informations sur le type d'export détecté
    """
    if export_type == 'samourai':
        return {
            'type': 'Type Samouraï',
            'icon': '⚔️',
            'description': 'Export Nota Bene classique',
            'colonnes': ['Reference', 'Nom', 'Prénom']
        }
    elif export_type == 'pirate':
        return {
            'type': 'Type Pirate',
            'icon': '🏴‍☠️', 
            'description': 'Export Ulule/Plateforme (priorité facturation/livraison)',
            'colonnes': ['#', 'Facturation/Livraison puis Nom complet']
        }
    else:
        return {
            'type': 'Type inconnu',
            'icon': '❓',
            'description': 'Structure non reconnue',
            'colonnes': list(df.columns)
        }

st.title("🔄 Fusion de Noms pour Remerciements")
st.markdown("*Application adaptative qui détecte automatiquement le type d'export*")

# Interface utilisateur
col1, col2 = st.columns(2)

with col1:
    st.subheader("📁 Fichiers requis")
    file_rem = st.file_uploader(
        "Fichier de remerciements (Excel)", 
        type=["xls", "xlsx"], 
        help="Export des contributeurs (Nota Bene, Ulule, etc.)"
    )
    file_changes = st.file_uploader(
        "Fichier de changements (Excel)", 
        type=["xls", "xlsx"], 
        help="Fichier avec les modifications à appliquer"
    )

with col2:
    st.subheader("📝 Configuration")
    file_name = st.text_input(
        "Nom du fichier Word à générer", 
        placeholder="remerciements_final", 
        help="Sans extension .docx"
    )
    
    # Options avancées
    with st.expander("🔧 Options avancées"):
        remove_duplicates = st.checkbox("Supprimer les doublons", value=True)
        show_preview = st.checkbox("Afficher un aperçu des données", value=False)
        case_sensitive = st.checkbox("Tri sensible à la casse", value=False)
        show_debug = st.checkbox("Afficher les informations de débogage", value=False)

if file_rem and file_changes and file_name:
    try:
        # Charger les fichiers
        with st.spinner("Chargement des fichiers..."):
            df_rem = pd.read_excel(file_rem)
            df_changes = pd.read_excel(file_changes)

        # Validation des fichiers vides
        if df_rem.empty or df_changes.empty:
            st.error("❌ L'un des fichiers est vide.")
            st.stop()

        # Détecter le type d'export
        export_type = detect_export_type(df_rem)
        export_info = get_export_info(export_type, df_rem)
        
        # Afficher les informations sur le type détecté
        st.success(f"✅ Fichiers chargés avec succès!")
        
        # Informations sur le type d'export
        st.subheader(f"{export_info['icon']} Type d'export détecté")
        
        col_type1, col_type2 = st.columns([1, 2])
        with col_type1:
            st.info(f"**{export_info['type']}**")
        with col_type2:
            st.info(f"📝 {export_info['description']}")
        
        if show_debug:
            st.expander("🔍 Informations de débogage").write({
                'Type détecté': export_type,
                'Colonnes disponibles': list(df_rem.columns),
                'Colonnes utilisées': export_info['colonnes']
            })

        # Gestion des types non reconnus
        if export_type == 'unknown':
            st.error("❌ **Type d'export non reconnu**")
            st.write("**Colonnes trouvées dans votre fichier:**")
            st.write(list(df_rem.columns))
            st.write("**Colonnes attendues:**")
            st.write("- **Type Samouraï**: Reference, Nom, Prénom")
            st.write("- **Type Pirate**: #, Nom complet (ou Prénom de facturation + Nom de facturation)")
            st.stop()

        # Standardiser le DataFrame
        with st.spinner("Standardisation des données..."):
            df_rem_std = standardize_dataframe(df_rem, export_type)
        
        # Afficher les statistiques
        col_info1, col_info2, col_info3 = st.columns(3)
        with col_info1:
            st.metric("📊 Contributeurs", len(df_rem_std))
        with col_info2:
            st.metric("🔄 Changements", len(df_changes))
        with col_info3:
            st.metric("📈 Type", export_info['type'])

        # Aperçu des données si demandé
        if show_preview:
            st.subheader("🔍 Aperçu des données standardisées")
            tab1, tab2 = st.tabs(["Données standardisées", "Changements"])
            
            with tab1:
                st.dataframe(df_rem_std[['Reference', 'Nom', 'Prénom']].head(10), use_container_width=True)
            with tab2:
                st.dataframe(df_changes.head(10), use_container_width=True)

        # Traitement des changements avec gestion des commentaires
        # Nettoyer et filtrer les changements
        df_changes = df_changes.dropna(subset=['Commande'])
        df_changes = df_changes[df_changes['Commande'].astype(str).str.contains('^[#]?[0-9a-zA-Z]+', regex=True)]
        
        # Vérifier les colonnes nécessaires dans le fichier changements
        required_cols_changes = ['Commande', 'A supprimer des pages Remerciements', 'A faire apparaitre sur les pages Remerciements']
        missing_cols_changes = [col for col in required_cols_changes if col not in df_changes.columns]
        
        if missing_cols_changes:
            st.error(f"❌ Colonnes manquantes dans le fichier changements: {missing_cols_changes}")
            st.stop()
        
        # Supprimer le symbole "#" dans les références de commandes
        df_changes['Commande'] = df_changes['Commande'].astype(str).str.replace('#', '', regex=False)
        
        # Nettoyer les espaces dans les colonnes de changements
        df_changes['A supprimer des pages Remerciements'] = df_changes['A supprimer des pages Remerciements'].astype(str).str.strip()
        df_changes['A faire apparaitre sur les pages Remerciements'] = df_changes['A faire apparaitre sur les pages Remerciements'].astype(str).str.strip()
        
        # Gérer la colonne Commentaire si elle existe
        exclude_from_thanks = set()
        if 'Commentaire' in df_changes.columns:
            # Identifier les personnes qui ne souhaitent pas apparaître
            df_changes['Commentaire'] = df_changes['Commentaire'].fillna('').astype(str).str.strip()
            
            # Conditions pour exclure quelqu'un
            exclude_conditions = df_changes['Commentaire'].str.lower().str.contains(
                'ne souhaite pas apparaître|ne souhaite pas apparaitre|ne pas apparaître|ne pas apparaitre', 
                regex=True, na=False
            )
            
            exclude_rows = df_changes[exclude_conditions]
            exclude_from_thanks = set(exclude_rows['Commande'].astype(str))
            
            if len(exclude_from_thanks) > 0:
                st.subheader("🚫 Exclusions demandées")
                st.info(f"**{len(exclude_from_thanks)} personne(s) ne souhaitent pas apparaître dans les remerciements**")
                
                # Afficher les exclusions
                exclude_df = exclude_rows[['Commande', 'A supprimer des pages Remerciements', 'Commentaire']].copy()
                exclude_df.columns = ['Référence', 'Nom à exclure', 'Motif']
                st.dataframe(exclude_df, use_container_width=True)

        # Vérification des correspondances
        rem_references = set(df_rem_std['Reference'].astype(str))
        change_references = set(df_changes['Commande'].astype(str))
        
        matches = change_references.intersection(rem_references)
        no_matches = change_references - rem_references
        
        st.subheader("🔍 Analyse des correspondances")
        col_match1, col_match2, col_match3 = st.columns(3)
        
        with col_match1:
            st.metric("✅ Références trouvées", len(matches))
        with col_match2:
            st.metric("❌ Références non trouvées", len(no_matches))
        with col_match3:
            st.metric("📊 Taux de correspondance", f"{(len(matches)/len(change_references)*100 if change_references else 0):.1f}%")
        
        if no_matches:
            st.warning(f"⚠️ **Références non trouvées:** {list(no_matches)}")

        # Créer le dictionnaire de remplacements (en excluant ceux qui ne veulent pas apparaître)
        replacements = {}
        
        for _, row in df_changes.iterrows():
            commande = str(row['Commande'])
            
            # Si la personne ne veut pas apparaître, on l'exclut complètement
            if commande in exclude_from_thanks:
                continue
                
            # Si il y a un remplacement spécifié et qu'il n'est pas vide
            nouveau_nom = row['A faire apparaitre sur les pages Remerciements']
            if pd.notna(nouveau_nom) and nouveau_nom.strip():
                replacements[commande] = nouveau_nom.strip()

        # Appliquer les remplacements
        df_rem_std['Nom Complet'] = df_rem_std.apply(
            lambda row: replacements.get(str(row['Reference']), f"{row['Prénom']} {row['Nom']}"), 
            axis=1
        )
        
        # Exclure complètement les personnes qui ne souhaitent pas apparaître
        if exclude_from_thanks:
            before_exclusion = len(df_rem_std)
            df_rem_std = df_rem_std[~df_rem_std['Reference'].astype(str).isin(exclude_from_thanks)]
            after_exclusion = len(df_rem_std)
            
            if before_exclusion > after_exclusion:
                st.success(f"✅ {before_exclusion - after_exclusion} personne(s) exclue(s) des remerciements comme demandé")

        # Nettoyer les espaces superflus
        df_rem_std['Nom Complet'] = df_rem_std['Nom Complet'].str.strip()
        df_rem_std['Nom Complet'] = df_rem_std['Nom Complet'].str.replace(r'\s+', ' ', regex=True)

        # Détection et gestion des doublons
        duplicates = df_rem_std['Nom Complet'][df_rem_std['Nom Complet'].duplicated(keep=False)]
        
        if not duplicates.empty:
            st.subheader("⚠️ Doublons détectés")
            st.warning(f"Nombre de doublons détectés : {len(duplicates)}")
            
            duplicate_df = df_rem_std[df_rem_std['Nom Complet'].isin(duplicates)][['Reference', 'Nom Complet']].sort_values('Nom Complet')
            st.dataframe(duplicate_df, use_container_width=True)
            
            if remove_duplicates:
                df_rem_std = df_rem_std.drop_duplicates(subset=['Nom Complet'])
                st.success(f"✅ Doublons supprimés automatiquement")

        # Fonction de tri (identique au code original)
        def normalize_sort_key(name):
            if pd.isna(name) or not name.strip():
                return 'zzz'
            
            prenom = name.split()[0] if name.split() else name
            normalized = unicodedata.normalize('NFD', prenom)
            no_accents = ''.join(c for c in normalized if unicodedata.category(c) != 'Mn')
            return no_accents.upper() if not case_sensitive else no_accents

        # Trier les noms
        with st.spinner("Tri des noms..."):
            sorted_names = sorted(
                df_rem_std['Nom Complet'].dropna().unique(),
                key=normalize_sort_key
            )
            sorted_names = [name.strip() for name in sorted_names if name.strip()]

        # Statistiques finales
        st.subheader("📈 Statistiques finales")
        col_stat1, col_stat2, col_stat3, col_stat4, col_stat5 = st.columns(5)
        
        with col_stat1:
            st.metric("👥 Total contributeurs", len(df_rem_std))
        with col_stat2:
            st.metric("🔄 Changements appliqués", len(replacements))
        with col_stat3:
            st.metric("🚫 Exclusions", len(exclude_from_thanks))
        with col_stat4:
            st.metric("📝 Noms finaux", len(sorted_names))
        with col_stat5:
            st.metric("🗑️ Doublons supprimés", len(duplicates) - len(df_rem_std[df_rem_std['Nom Complet'].duplicated()]) if remove_duplicates else 0)

        # Aperçu des premiers noms triés
        if st.checkbox("📋 Aperçu des 20 premiers noms triés"):
            st.write("**Premiers noms dans l'ordre alphabétique:**")
            preview_names = sorted_names[:20]
            for i, name in enumerate(preview_names, 1):
                st.write(f"{i:2d}. {name}")
            if len(sorted_names) > 20:
                st.write(f"... et {len(sorted_names) - 20} autres noms")

        # Création du document Word
        with st.spinner("Génération du document Word..."):
            doc = Document()
            doc.add_heading("Remerciements", 0)
            
            names_text = ", ".join(sorted_names)
            doc.add_paragraph(names_text)
            
            # Métadonnées
            doc.add_paragraph()
            doc.add_paragraph(f"Total des contributeurs: {len(sorted_names)}")
            doc.add_paragraph(f"Type d'export: {export_info['type']}")
            if exclude_from_thanks:
                doc.add_paragraph(f"Exclusions appliquées: {len(exclude_from_thanks)}")
            doc.add_paragraph(f"Document généré le: {pd.Timestamp.now().strftime('%d/%m/%Y à %H:%M')}")

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

        # Bouton de téléchargement
        word_file_name = f"{file_name}.docx"
        
        st.download_button(
            label=f"📥 Télécharger le fichier de remerciements (Word)",
            data=buffer,
            file_name=word_file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            help=f"Télécharger '{word_file_name}' avec {len(sorted_names)} noms triés"
        )
        
        st.success(f"✅ **Document '{word_file_name}' prêt au téléchargement!**")
        st.balloons()

    except Exception as e:
        st.error(f"❌ **Erreur lors du traitement:** {str(e)}")
        if show_debug:
            st.exception(e)

else:
    st.info("👆 **Veuillez téléverser les deux fichiers Excel et saisir un nom de fichier pour commencer.**")
    